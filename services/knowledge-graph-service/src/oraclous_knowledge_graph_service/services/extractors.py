"""Document → text extractors (services layer) — a pluggable per-content-type registry.

Reshaped from legacy `develop@84152635 knowledge-graph-builder/app/services/{document_processor,
pdf_extractor,vision_extractor,md_extractor}.py`. The legacy code was four free-standing modules
that each took a file *path*; the new ingest spine carries *bytes* and dispatches on a resolved
`source_type`/extension. This module is the dispatcher: each first-class document content-type is a
registered `DocumentExtractor` (a `(data, filename, source_type) -> (text, metadata)` callable),
looked up by the resolved *kind*. New types plug in by registering an extractor — no change to the
`extract_text` call site (`IngestionService`) or the `/upload` validation (`source_type_for`).

Built-in extractors (all feed the SAME text ingest path — document → text → chunk → embed → write):
  - text : UTF-8 decode (zero deps).
  - md   : UTF-8 decode; text is the RAW markdown plus a structured {title, sections, hierarchy}
           sidecar from the stdlib-`re` heading parser (lifted verbatim — no markdown lib).
  - pdf  : `pypdf` per-page text ("[Page N]\\n…"); richer table extraction via `pdfplumber` when
           installed (appended as "[Table]\\n…"), gracefully skipped otherwise. No OCR.
  - docx : `python-docx` paragraphs + table cells.
  - image: vision (image → entities/relationships prose) via the live LLM seam, the SAME OpenRouter
           config as `KGS_EXTRACTOR` (see `vision_extractor`). Needs a configured LLM — fail-closed.

Structured kinds (csv/tsv/json/jsonl) and code are recognised by `_kind` so `/upload` validates +
stores the right `source_type`, but `extract_text` never handles them — the worker routes those to
`StructuredIngestionService` / `CodeIngestionService` instead.

Raises plain `ExtractionError` (the route layer maps it to HTTP) — no FastAPI coupling in services.
"""

from __future__ import annotations

import io
import logging
import re
from collections.abc import Callable
from typing import Any

from oraclous_knowledge_graph_service.core.config import Settings, get_settings
from oraclous_knowledge_graph_service.services.vision_extractor import (
    is_image_ext,
    make_vision_extractor,
)

logger = logging.getLogger(__name__)

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$", re.MULTILINE)
_TEXT_EXTS = {"txt", "text", "log", ""}
_MD_EXTS = {"md", "markdown", "mdown"}
_CODE_EXTS = {"py", "ts", "tsx", "js", "jsx", "go", "java", "zip"}


class ExtractionError(Exception):
    """Extraction failed (unsupported type, corrupt file, empty text, or vision off). HTTP 422."""


# A document extractor: bytes (+ optional filename/source_type) -> (plain_text, metadata).
DocumentExtractor = Callable[..., "tuple[str, dict[str, Any]]"]


def _ext(filename: str | None) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def _kind(source_type: str | None, filename: str | None) -> str:
    """Resolve a source_type/filename to a canonical *kind* (the registry key / routing key)."""
    st = (source_type or "").lower()
    if st in {"text", "txt", "md", "markdown", "pdf", "docx", "doc", "image"}:
        if st in {"md", "markdown"}:
            return "md"
        if st in {"docx", "doc"}:
            return "docx"
        if st in {"text", "txt"}:
            return "text"
        return st
    if st in {"csv", "tsv", "json", "jsonl"}:
        return st
    if st == "code":
        return "code"
    ext = _ext(filename)
    if ext == "pdf":
        return "pdf"
    if ext in {"docx", "doc"}:
        return "docx"
    if ext in {"csv", "tsv", "json", "jsonl"}:
        return ext
    if ext in _CODE_EXTS:
        return "code"
    if ext in _MD_EXTS:
        return "md"
    if is_image_ext(ext):
        return "image"
    if ext in _TEXT_EXTS:
        return "text"
    raise ExtractionError(f"unsupported source type: {source_type or filename!r}")


# --- markdown (stdlib heading parser, lifted verbatim) -----------------------
def _build_hierarchy(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    roots: list[dict[str, Any]] = []
    stack: list[dict[str, Any]] = []
    for sec in sections:
        node = {"level": sec["level"], "heading": sec["heading"], "children": []}
        while stack and stack[-1]["level"] >= node["level"]:
            stack.pop()
        if stack:
            stack[-1]["children"].append(node)
        else:
            roots.append(node)
        stack.append(node)
    return roots


def _extract_markdown(text: str, fallback_title: str) -> dict[str, Any]:
    matches = list(_HEADING_RE.finditer(text))
    sections: list[dict[str, Any]] = []
    for i, match in enumerate(matches):
        level = len(match.group(1))
        heading = match.group(2).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        sections.append({"level": level, "heading": heading, "content": text[start:end].strip()})
    title = next((s["heading"] for s in sections if s["level"] == 1), fallback_title)
    return {"title": title, "sections": sections, "hierarchy": _build_hierarchy(sections)}


# --- PDF (pypdf text + optional pdfplumber tables) ---------------------------
def _extract_pdf_tables(data: bytes) -> list[str]:
    """Per-table "[Table]\\n<tab-rows>" blocks via pdfplumber; [] if pdfplumber absent or fails.

    pdfplumber is an optional richer-extraction dep — table mining is best-effort and never sinks a
    PDF whose text already parsed (legacy pdf_extractor behaviour).
    """
    try:
        import pdfplumber
    except ImportError:
        logger.debug("pdfplumber not installed — PDF table extraction skipped")
        return []
    blocks: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables() or []:
                    rows = [
                        "\t".join("" if cell is None else str(cell) for cell in row)
                        for row in table
                        if row
                    ]
                    rows = [r for r in rows if r.strip()]
                    if rows:
                        blocks.append("[Table]\n" + "\n".join(rows))
    except Exception as exc:  # pragma: no cover - corrupt-pdf table path
        logger.warning("PDF table extraction failed: %s", exc)
        return []
    return blocks


def _extract_pdf(
    *, data: bytes, filename: str | None = None, **_: Any
) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dep is declared
        raise ExtractionError("pypdf is required to extract PDF files") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"could not read PDF: {exc}") from exc
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            parts.append(f"[Page {i + 1}]\n{page_text}")
    tables = _extract_pdf_tables(data)
    parts.extend(tables)
    text = "\n\n".join(parts)
    metadata: dict[str, Any] = {
        "kind": "pdf",
        "filename": filename,
        "page_count": len(reader.pages),
        "has_tables": bool(tables),
    }
    return text, metadata


def _extract_docx(
    *, data: bytes, filename: str | None = None, **_: Any
) -> tuple[str, dict[str, Any]]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dep is declared
        raise ExtractionError("python-docx is required to extract DOCX files") from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"could not read DOCX: {exc}") from exc
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    has_tables = False
    for table in document.tables:
        rows = ["\t".join(c.text.strip() for c in row.cells) for row in table.rows]
        rows = [r for r in rows if r.strip()]
        if rows:
            has_tables = True
            parts.append("[Table]\n" + "\n".join(rows))
    return "\n\n".join(parts), {"kind": "docx", "filename": filename, "has_tables": has_tables}


def _extract_text(
    *, data: bytes, filename: str | None = None, **_: Any
) -> tuple[str, dict[str, Any]]:
    text = data.decode("utf-8", errors="replace")
    return text, {"kind": "text", "filename": filename}


def _extract_md(
    *, data: bytes, filename: str | None = None, **_: Any
) -> tuple[str, dict[str, Any]]:
    text = data.decode("utf-8", errors="replace")  # md text is the RAW markdown
    structured = _extract_markdown(text, fallback_title=filename or "")
    return text, {"kind": "md", "filename": filename, "structured": structured}


def _extract_image(
    *, data: bytes, filename: str | None = None, settings: Settings | None = None, **_: Any
) -> tuple[str, dict[str, Any]]:
    """Image → prose via the live LLM vision seam; fail-closed when no LLM is configured."""
    vision = make_vision_extractor(settings or get_settings())
    if vision is None:
        raise ExtractionError(
            "image extraction requires the LLM seam (set KGS_EXTRACTOR=openai + KGS_OPENAI_API_KEY)"
        )
    return vision.extract(data=data, filename=filename)


# The pluggable registry: kind -> extractor. New first-class types register here.
_REGISTRY: dict[str, DocumentExtractor] = {
    "text": _extract_text,
    "md": _extract_md,
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "image": _extract_image,
}


def register_extractor(kind: str, extractor: DocumentExtractor) -> None:
    """Register (or override) the extractor for a kind — the plug-in seam."""
    _REGISTRY[kind] = extractor


def supported_kinds() -> tuple[str, ...]:
    """The document kinds with a registered text extractor (excludes structured/code routing)."""
    return tuple(_REGISTRY)


def source_type_for(filename: str | None) -> str:
    """Resolve + validate the source_type for an upload (raises ExtractionError if unsupported)."""
    return _kind(None, filename)


def extract_text(
    *,
    data: bytes,
    filename: str | None = None,
    source_type: str | None = None,
    settings: Settings | None = None,
) -> tuple[str, dict[str, Any]]:
    """Return (plain_text, metadata) by dispatching on content-type. Raises ExtractionError.

    Looks up the resolved *kind* in the pluggable registry and delegates. `settings` flows to
    extractors that need the LLM seam (image/vision); it defaults to the process settings.
    """
    kind = _kind(source_type, filename)
    extractor = _REGISTRY.get(kind)
    if extractor is None:
        # csv/tsv/json/jsonl/code resolve to a kind but are routed elsewhere by the worker; a direct
        # extract_text on one is a programming error, surfaced (never silently emptied).
        raise ExtractionError(f"no text extractor for kind {kind!r} (routed to a structured path)")
    text, metadata = extractor(
        data=data, filename=filename, source_type=source_type, settings=settings
    )
    text = text.strip()
    if not text:
        raise ExtractionError("no extractable text found")
    return text, metadata
